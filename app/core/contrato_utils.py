import uuid
from typing import Optional
from uuid import UUID
from fastapi import Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from datetime import datetime
import locale

from sqlalchemy.orm import joinedload
from starlette import status
from starlette.responses import JSONResponse
from sqlalchemy.sql import delete
from app.Enum.StatusContrato import StatusContrato
from app.Enum.StatusCobranca import StatusCobranca
from app.connection.database import get_db
from app.core.anexo_utils import base64_to_bytes, bytes_to_base64
from app.core.auth_utils import verificar_token
from app.core.log_utils import limpar_dict_para_json
from app.models.ContratoModel import ContratoModel
from app.models.ClienteModel import ClienteModel
from app.models.AnexoModel import AnexoModel
from app.models.EnderecoModel import EnderecoModel
from app.models.LogModel import LogModel
from app.models.ParcelamentoModel import ParcelamentoModel
from app.models.PlanoModel import PlanoModel
from sqlalchemy import func, and_, or_, cast, String

from app.models.ServicoModel import ServicoModel
from app.models.VendedorModel import VendedorModel
from app.schemas.AnexoSchema import AnexoRequest
from app.schemas.ClienteSchema import  ClienteContratoResponse
from app.schemas.ContratoSchema import ContratoRequest, ContratoResponse, ContratoResponseShort
from app.schemas.EnderecoSchema import EnderecoRequest
from app.schemas.ParcelamentoSchema import ParcelamentoResponse
from app.schemas.PlanoSchema import PlanoUpdate, PlanoRequest, PlanoServicoResponse
from sqlalchemy.future import select

from app.schemas.ServicoSchema import ServicoList
from app.schemas.VendedorSchema import VendedorContratoResponse
from datetime import date
from dateutil.relativedelta import relativedelta

from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from num2words import num2words
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


async def criar(form_data: ContratoRequest,
                     db: AsyncSession = Depends(get_db), user_id: str = Depends(verificar_token)):

    queryPlano = select(ContratoModel).where(
        and_(
                ContratoModel.cliente_id == form_data.cliente_id,
                ContratoModel.ativo == True,
                ContratoModel.plano_id == form_data.plano_id,
        )
    )
    resultPlano = await db.execute(queryPlano)
    contratoPlano = resultPlano.scalar_one_or_none()
    if contratoPlano:
        raise HTTPException(status_code=400, detail=f"Já existe um contrato ativo para este cliente com o mesmo plano selecionado.")

    form_data.id = uuid.uuid4()

    if not form_data.cliente_id:
        raise HTTPException(status_code=400, detail=f"Cliente inválido, tente novamente!")

    queryClienteId = select(ClienteModel).where(
        and_(
            ClienteModel.id == form_data.cliente_id,
            ClienteModel.ativo == True
        )
    )
    resultClienteId = await db.execute(queryClienteId)
    contratoClienteId = resultClienteId.scalar_one_or_none()
    if not contratoClienteId:
        raise HTTPException(status_code=400, detail=f"Cliente não localizado ou inativo, verifique os dados e tente novamente.")


    if not form_data.plano_id:
        raise HTTPException(status_code=400, detail=f"Plano inválido, tente novamente!")

    queryPlanoId = select(PlanoModel).where(
        and_(
            PlanoModel.id == form_data.plano_id,
            PlanoModel.ativo == True
        )
    )
    resultPlanoId = await db.execute(queryPlanoId)
    contratoPlanoId = resultPlanoId.scalar_one_or_none()
    if not contratoPlanoId:
        raise HTTPException(status_code=400,
                            detail=f"Plano não localizado ou inativo, verifique os dados e tente novamente.")


    if form_data.vendedor_id is not None:
        queryVendedorId = select(VendedorModel).where(
            and_(
                VendedorModel.id == form_data.vendedor_id,
                VendedorModel.ativo == True
            )
        )
        resultVendedorId = await db.execute(queryVendedorId)
        contratoVendedorId = resultVendedorId.scalar_one_or_none()
        if not contratoVendedorId:
            raise HTTPException(status_code=400,
                                detail=f"Vendedor não localizado ou inativo, verifique os dados e tente novamente.")



    if not form_data.parcelamento.data_inicio:
        raise HTTPException(status_code=400, detail=f"Por favor, escolha a data de início.")


    if form_data.parcelamento.valor_total <= 0 or form_data.parcelamento.valor_parcela <= 0:
        raise HTTPException(
            status_code=400,
            detail="O valor total e o valor da parcela devem ser maiores que zero."
        )




    anexos_id = []
    if form_data.anexos_list is not None:
        for anexo in form_data.anexos_list:
            image_bytes = base64_to_bytes(anexo.base64)
            novo_anexo = AnexoModel(
                id=uuid.uuid4(),
                base64=image_bytes,
                image=anexo.image,
                descricao=anexo.descricao,
                nome=anexo.nome,
                tipo=anexo.tipo,
                created_by=uuid.UUID(user_id)
            )
            db.add(novo_anexo)
            await db.flush()
            anexos_id.append(novo_anexo.id)

    parcelamento_id = uuid.uuid4()
    novo_parcelamento = ParcelamentoModel(
        id=parcelamento_id,
        contrato_id=form_data.id,
        data_inicio= form_data.parcelamento.data_inicio,
        data_fim= form_data.parcelamento.data_fim,
        data_vigencia= form_data.parcelamento.data_vigencia,
        meio_pagamento= form_data.parcelamento.meio_pagamento,
        valor_total= form_data.parcelamento.valor_total,
        valor_parcela= form_data.parcelamento.valor_parcela,
        valor_entrada= form_data.parcelamento.valor_entrada,
        qtd_parcela= form_data.parcelamento.qtd_parcela,
        avista= form_data.parcelamento.avista,
        taxa_juros= form_data.parcelamento.taxa_juros,
        data_ultimo_pagamento= form_data.parcelamento.data_ultimo_pagamento,
        qtd_parcelas_pagas= form_data.parcelamento.qtd_parcelas_pagas,
        ativo=True,
        created_by=uuid.UUID(user_id)
    )
    db.add(novo_parcelamento)
    await db.flush()

    novo_contrato = ContratoModel(
        id=form_data.id,
        parcelamento_id=parcelamento_id,
        cliente_assinatura_id=None,
        responsavel_assinatura_id=None,
        cliente_id=form_data.cliente_id,
        vendedor_id=form_data.vendedor_id,
        plano_id=form_data.plano_id,
        anexos_list_id=anexos_id,
        servicos_list_id=None,
        nome=contratoClienteId.nome,
        documento=contratoClienteId.documento,
        status_cobranca=StatusCobranca.EM_DIA.value,
        status_contrato=StatusContrato.INICIADO.value,
        ativo=True,
        created_by=uuid.UUID(user_id)
    )
    db.add(novo_contrato)
    await db.flush()

    dados_antigos = None
    dados_novos =  limpar_dict_para_json(form_data)

    log = LogModel(
        tabela_afetada="contratos",
        operacao="CREATE",
        registro_id=form_data.id,
        dados_antes=dados_antigos,
        dados_depois=dados_novos,
        usuario_id=uuid.UUID(user_id)
    )

    db.add(log)

    await db.commit()
    await db.refresh(novo_contrato)

    return JSONResponse(
        content={"detail": "Contrato criado com sucesso"},
        media_type="application/json; charset=utf-8"
    )

async def listar(
    pagina: int = Query(1, ge=1),
    items: int = Query(10, ge=1, le=100),
    filtro: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db)
):
    offset = (pagina - 1) * items

    where_clause = [ContratoModel.deleted_at == None]

    total_query = select(func.count(ContratoModel.id))

    if filtro:
        filtro_str = f"%{filtro.lower()}%"
        where_clause.append(
            or_(
                func.lower(cast(ContratoModel.numero, String)).ilike(filtro_str),
                func.lower(ContratoModel.status_contrato).ilike(filtro_str),
                func.lower(ClienteModel.nome).ilike(filtro_str),
                func.lower(cast(ClienteModel.documento, String)).ilike(filtro_str)
            )
        )
        total_query = total_query.join(ClienteModel, ContratoModel.cliente_id == ClienteModel.id)


    total_query = total_query.where(*where_clause)
    total_result = await db.execute(total_query)
    total_items = total_result.scalar()
    total_paginas = (total_items + items - 1) // items if total_items > 0 else 0

    query = select(ContratoModel).options(
        joinedload(ContratoModel.parcelamento),
        joinedload(ContratoModel.responsavel_assinatura),
        joinedload(ContratoModel.cliente_assinatura),
        joinedload(ContratoModel.cliente),
        joinedload(ContratoModel.vendedor),
        joinedload(ContratoModel.plano),
    )

    if filtro:
        query = query.join(ClienteModel, ContratoModel.cliente_id == ClienteModel.id)

    query = query.where(*where_clause).offset(offset).limit(items)

    result = await db.execute(query)
    contratos = result.scalars().unique().all()

    contratos_list = []
    for contrato in contratos:
        parcelamento = None
        if contrato.parcelamento:
            parcelamento = ParcelamentoResponse(
                id=contrato.parcelamento.id,
                data_inicio=contrato.parcelamento.data_inicio,
                data_fim=contrato.parcelamento.data_fim,
                data_vigencia=contrato.parcelamento.data_vigencia,
                meio_pagamento=contrato.parcelamento.meio_pagamento,
                valor_total=contrato.parcelamento.valor_total,
                valor_parcela=contrato.parcelamento.valor_parcela,
                valor_entrada=contrato.parcelamento.valor_entrada,
                qtd_parcela=contrato.parcelamento.qtd_parcela,
                avista=contrato.parcelamento.avista,
                taxa_juros=contrato.parcelamento.taxa_juros,
                data_ultimo_pagamento=contrato.parcelamento.data_ultimo_pagamento,
                qtd_parcelas_pagas=contrato.parcelamento.qtd_parcelas_pagas,
                ativo=contrato.parcelamento.ativo,

                created_at = contrato.parcelamento.created_at,
                updated_at = contrato.parcelamento.updated_at,
                deleted_at=contrato.parcelamento.deleted_at,
                created_by=contrato.parcelamento.created_by,
                updated_by=contrato.parcelamento.updated_by,
                deleted_by=contrato.parcelamento.deleted_by,
            )

        anexo_list = []
        if contrato.anexos_list_id:
            for anexo_id in contrato.anexos_list_id:
                query = select(AnexoModel).where(
                    and_(
                        AnexoModel.id == anexo_id
                    )
                )
                result = await db.execute(query)
                item = result.scalar_one_or_none()

                if item:
                    image_base64 = bytes_to_base64(item.base64, item.tipo)
                    anexo = AnexoRequest(
                        id=item.id,
                        image=item.image,
                        base64=image_base64,
                        descricao=item.descricao,
                        nome=item.nome,
                        tipo=item.tipo,
                    )
                    anexo_list.append(anexo)

        cliente_assinatura = None
        if contrato.cliente_assinatura:
            image_base64 = bytes_to_base64(contrato.cliente_assinatura.base64, contrato.cliente_assinatura.tipo)
            cliente_assinatura = AnexoRequest(
                id=contrato.cliente_assinatura.id,
                image=contrato.cliente_assinatura.image,
                base64=image_base64,
                descricao=contrato.cliente_assinatura.descricao,
                nome=contrato.cliente_assinatura.nome,
                tipo=contrato.cliente_assinatura.tipo,
            )

        responsavel_assinatura = None
        if contrato.responsavel_assinatura:
            image_base64 = bytes_to_base64(contrato.responsavel_assinatura.base64, contrato.responsavel_assinatura.tipo)
            responsavel_assinatura = AnexoRequest(
                id=contrato.responsavel_assinatura.id,
                image=contrato.responsavel_assinatura.image,
                base64=image_base64,
                descricao=contrato.responsavel_assinatura.descricao,
                nome=contrato.responsavel_assinatura.nome,
                tipo=contrato.responsavel_assinatura.tipo,
            )

        cliente = None
        if contrato.cliente:
            cliente = ClienteContratoResponse(
                id=contrato.cliente.id,
                nome=contrato.cliente.nome,
                documento=contrato.cliente.documento,
                email=contrato.cliente.email,
                telefone=contrato.cliente.telefone,
                grupo_segmento=contrato.cliente.grupo_segmento,
                ativo=contrato.cliente.ativo
            )

        vendedor = None
        if contrato.vendedor:
            vendedor = VendedorContratoResponse(
                id=contrato.vendedor.id,
                nome=contrato.vendedor.nome,
                cpf=contrato.vendedor.cpf,
                email=contrato.vendedor.email,
                telefone=contrato.vendedor.telefone,
                rg=contrato.vendedor.rg,
                ativo=contrato.vendedor.ativo
            )

        plano = None
        if contrato.plano:
            plano = PlanoRequest(
                id=contrato.plano.id,
                nome=contrato.plano.nome,
                descricao=contrato.plano.descricao,
                valor_mensal=contrato.plano.valor_mensal,
                numero_parcelas=contrato.plano.numero_parcelas,
                avista=contrato.plano.avista,
                periodo_vigencia=contrato.plano.periodo_vigencia,
                servicos_vinculados=contrato.plano.servicos_vinculados
            )

        response = ContratoResponseShort(
            id=contrato.id,
            numero=contrato.numero,
            ativo=contrato.ativo,
            status_cobranca=contrato.status_cobranca,
            status_contrato=contrato.status_contrato,
            parcelamento=parcelamento,
            plano=plano,
            cliente=cliente,
            created_at=contrato.created_at,
            updated_at=contrato.updated_at
        )
        contratos_list.append(response)

    return {
        "total_items": total_items,
        "total_paginas": total_paginas,
        "pagina_atual": pagina,
        "items": items,
        "offset": offset,
        "data": contratos_list
    }



async def por_id(id: uuid.UUID,
                   db: AsyncSession = Depends(get_db)):

    query = select(ContratoModel).options(
        joinedload(ContratoModel.parcelamento),
        joinedload(ContratoModel.responsavel_assinatura),
        joinedload(ContratoModel.cliente_assinatura),
        joinedload(ContratoModel.cliente),
        joinedload(ContratoModel.vendedor),
        joinedload(ContratoModel.plano),
    ).where(ContratoModel.id == id)

    result = await db.execute(query)
    contrato = result.scalar_one_or_none()
    if not contrato:
        raise HTTPException(status_code=400, detail="Contrato não localizado na base de dados.")

    parcelamento = None
    if contrato.parcelamento:
        parcelamento = ParcelamentoResponse(
            id=contrato.parcelamento.id,
            data_inicio=contrato.parcelamento.data_inicio,
            data_fim=contrato.parcelamento.data_fim,
            data_vigencia=contrato.parcelamento.data_vigencia,
            meio_pagamento=contrato.parcelamento.meio_pagamento,
            valor_total=contrato.parcelamento.valor_total,
            valor_parcela=contrato.parcelamento.valor_parcela,
            valor_entrada=contrato.parcelamento.valor_entrada,
            qtd_parcela=contrato.parcelamento.qtd_parcela,
            avista=contrato.parcelamento.avista,
            taxa_juros=contrato.parcelamento.taxa_juros,
            data_ultimo_pagamento=contrato.parcelamento.data_ultimo_pagamento,
            qtd_parcelas_pagas=contrato.parcelamento.qtd_parcelas_pagas,
            ativo=contrato.parcelamento.ativo,

            created_at=contrato.parcelamento.created_at,
            updated_at=contrato.parcelamento.updated_at,
            deleted_at=contrato.parcelamento.deleted_at,
            created_by=contrato.parcelamento.created_by,
            updated_by=contrato.parcelamento.updated_by,
            deleted_by=contrato.parcelamento.deleted_by,
        )

    anexo_list = []
    if contrato.anexos_list_id:
        for anexo_id in contrato.anexos_list_id:
            query = select(AnexoModel).where(
                and_(
                    AnexoModel.id == anexo_id
                )
            )
            result = await db.execute(query)
            item = result.scalar_one_or_none()

            if item:
                image_base64 = bytes_to_base64(item.base64, item.tipo)
                anexo = AnexoRequest(
                    id=item.id,
                    image=item.image,
                    base64=image_base64,
                    descricao=item.descricao,
                    nome=item.nome,
                    tipo=item.tipo,
                )
                anexo_list.append(anexo)

    cliente_assinatura = None
    if contrato.cliente_assinatura:
        image_base64 = bytes_to_base64(contrato.cliente_assinatura.base64, contrato.cliente_assinatura.tipo)
        cliente_assinatura = AnexoRequest(
            id=contrato.cliente_assinatura.id,
            image=contrato.cliente_assinatura.image,
            base64=image_base64,
            descricao=contrato.cliente_assinatura.descricao,
            nome=contrato.cliente_assinatura.nome,
            tipo=contrato.cliente_assinatura.tipo,
        )

    responsavel_assinatura = None
    if contrato.responsavel_assinatura:
        image_base64 = bytes_to_base64(contrato.responsavel_assinatura.base64, contrato.responsavel_assinatura.tipo)
        responsavel_assinatura = AnexoRequest(
            id=contrato.responsavel_assinatura.id,
            image=contrato.responsavel_assinatura.image,
            base64=image_base64,
            descricao=contrato.responsavel_assinatura.descricao,
            nome=contrato.responsavel_assinatura.nome,
            tipo=contrato.responsavel_assinatura.tipo,
        )

    cliente = None
    if contrato.cliente:
        endereco_cliente = None
        query = select(EnderecoModel).where(
            and_(
                EnderecoModel.id == contrato.cliente.endereco_id,
            )
        )

        result = await db.execute(query)
        item = result.scalar_one_or_none()
        if item:
            endereco_cliente = EnderecoRequest(
                id=item.id,
                cep=item.cep,
                rua=item.rua,
                numero=item.numero,
                bairro=item.bairro,
                complemento=item.complemento,
                cidade=item.cidade,
                uf=item.uf,
            )

        cliente = ClienteContratoResponse(
            id=contrato.cliente.id,
            nome=contrato.cliente.nome,
            documento=contrato.cliente.documento,
            email=contrato.cliente.email,
            telefone=contrato.cliente.telefone,
            grupo_segmento=contrato.cliente.grupo_segmento,
            ativo=contrato.cliente.ativo,
            endereco=endereco_cliente
        )

    vendedor = None
    if contrato.vendedor:
        vendedor = VendedorContratoResponse(
            id=contrato.vendedor.id,
            nome=contrato.vendedor.nome,
            cpf=contrato.vendedor.cpf,
            email=contrato.vendedor.email,
            telefone=contrato.vendedor.telefone,
            rg=contrato.vendedor.rg,
            ativo=contrato.vendedor.ativo
        )

    servicos_list = []
    if contrato.plano.servicos_vinculados:
        for servico in contrato.plano.servicos_vinculados:
            query = select(ServicoModel).where(
                and_(
                    ServicoModel.id == servico
                )
            )
            result = await db.execute(query)
            item = result.scalar_one_or_none()

            if item:
                servicoItem = ServicoList(
                    id=item.id,
                    nome=item.nome,
                    valor=item.valor
                )
                servicos_list.append(servicoItem)

    plano = None
    if contrato.plano:
        plano = PlanoServicoResponse(
            id=contrato.plano.id,
            nome=contrato.plano.nome,
            descricao=contrato.plano.descricao,
            valor_mensal=contrato.plano.valor_mensal,
            numero_parcelas=contrato.plano.numero_parcelas,
            ativo=contrato.plano.ativo,
            avista=contrato.plano.avista,
            periodo_vigencia=contrato.plano.periodo_vigencia,
            servicos_vinculados=servicos_list
        )

    response = ContratoResponse(
        id=contrato.id,
        numero=contrato.numero,
        nome=contrato.nome,
        documento=contrato.documento,
        ativo=contrato.ativo,
        status_cobranca=contrato.status_cobranca,
        status_contrato=contrato.status_contrato,
        parcelamento=parcelamento,
        cliente=cliente,
        vendedor=vendedor,
        plano=plano,
        anexos_list=anexo_list,
        responsavel_assinatura=responsavel_assinatura,
        cliente_assinatura=cliente_assinatura,
        created_at=contrato.created_at,
        updated_at=contrato.updated_at,
        deleted_at=contrato.deleted_at,
        created_by=contrato.created_by,
        updated_by=contrato.updated_by,
        deleted_by=contrato.deleted_by,
    )

    return response


async def atualizar(id: uuid.UUID, form_data: PlanoUpdate,
                   db: AsyncSession = Depends(get_db), user_id: str = Depends(verificar_token)):
    queryContrato = select(ContratoModel).where(ContratoModel.id == id)
    resultContrato = await db.execute(queryContrato)
    contrato_existente = resultContrato.scalar_one_or_none()

    dados_antes = limpar_dict_para_json(contrato_existente)

    if not contrato_existente:
        raise HTTPException(status_code=404, detail="Contrato não encontrado.")

    queryPlanoDuplicado = select(ContratoModel).where(
        and_(
            ContratoModel.cliente_id == form_data.cliente_id,
            ContratoModel.plano_id == form_data.plano_id,
            ContratoModel.ativo == True,
            ContratoModel.id != id
        )
    )
    resultPlanoDuplicado = await db.execute(queryPlanoDuplicado)
    if resultPlanoDuplicado.scalar_one_or_none():
        raise HTTPException(
            status_code=400,
            detail="Já existe um contrato ativo para este cliente com o mesmo plano."
        )

    queryCliente = select(ClienteModel).where(
        and_(
            ClienteModel.id == form_data.cliente_id,
            ClienteModel.ativo == True
        )
    )
    resultCliente = await db.execute(queryCliente)
    cliente = resultCliente.scalar_one_or_none()
    if not cliente:
        raise HTTPException(status_code=400, detail="Cliente não localizado ou inativo.")

    queryPlano = select(PlanoModel).where(
        and_(
            PlanoModel.id == form_data.plano_id,
            PlanoModel.ativo == True
        )
    )
    resultPlano = await db.execute(queryPlano)
    plano = resultPlano.scalar_one_or_none()
    if not plano:
        raise HTTPException(status_code=400, detail="Plano não localizado ou inativo.")

    if form_data.vendedor_id:
        queryVendedor = select(VendedorModel).where(
            and_(
                VendedorModel.id == form_data.vendedor_id,
                VendedorModel.ativo == True
            )
        )
        resultVendedor = await db.execute(queryVendedor)
        vendedor = resultVendedor.scalar_one_or_none()
        if not vendedor:
            raise HTTPException(status_code=400, detail="Vendedor não localizado ou inativo.")

    if not form_data.parcelamento.data_inicio:
        raise HTTPException(status_code=400, detail="Data de início obrigatória.")

    if form_data.parcelamento.valor_total <= 0 or form_data.parcelamento.valor_parcela <= 0:
        raise HTTPException(
            status_code=400,
            detail="O valor total e o valor da parcela devem ser maiores que zero."
        )

    anexos_antigos_ids = set(contrato_existente.anexos_list_id or [])
    anexos_novos_ids = set()
    anexos_final_ids = []

    if form_data.anexos_list:
        for anexo in form_data.anexos_list:
            if anexo.id:
                query_anexo_existente = select(AnexoModel).where(AnexoModel.id == anexo.id)
                result_anexo = await db.execute(query_anexo_existente)
                anexo_existente = result_anexo.scalar_one_or_none()

                if anexo_existente:
                    anexo_existente.base64 = base64_to_bytes(anexo.base64)
                    anexo_existente.image = anexo.image
                    anexo_existente.descricao = anexo.descricao
                    anexo_existente.nome = anexo.nome
                    anexo_existente.tipo = anexo.tipo
                    anexo_existente.updated_by = uuid.UUID(user_id)
                    anexo_existente.updated_at = datetime.utcnow()

                    await db.merge(anexo_existente)

                    anexos_novos_ids.add(anexo.id)
                    anexos_final_ids.append(anexo.id)
            else:
                novo_item_id = uuid.uuid4()
                image_bytes = base64_to_bytes(anexo.base64)
                novo_anexo = AnexoModel(
                    id=novo_item_id,
                    base64=image_bytes,
                    image=anexo.image,
                    descricao=anexo.descricao,
                    nome=anexo.nome,
                    tipo=anexo.tipo,
                    created_by=uuid.UUID(user_id)
                )
                db.add(novo_anexo)
                await db.flush()
                anexos_novos_ids.add(novo_item_id)
                anexos_final_ids.append(novo_item_id)

    anexos_para_apagar = anexos_antigos_ids - anexos_novos_ids
    if anexos_para_apagar:
        delete_query = delete(AnexoModel).where(AnexoModel.id.in_(anexos_para_apagar))
        await db.execute(delete_query)

    parcelamento_id = None
    if form_data.parcelamento.id:
        query_parcelamento_existente = select(ParcelamentoModel).where(ParcelamentoModel.id == form_data.parcelamento.id)
        result_parcelamento = await db.execute(query_parcelamento_existente)
        parcelamento_existente = result_parcelamento.scalar_one_or_none()


        if parcelamento_existente:
            parcelamento_id = form_data.parcelamento.id
            if parcelamento_existente.data_inicio is not None:
                parcelamento_existente.data_inicio = form_data.parcelamento.data_inicio
            if parcelamento_existente.data_fim is not None:
                parcelamento_existente.data_fim = form_data.parcelamento.data_fim
            if parcelamento_existente.data_vigencia is not None:
                parcelamento_existente.data_vigencia = form_data.parcelamento.data_vigencia
            if parcelamento_existente.meio_pagamento is not None:
                parcelamento_existente.meio_pagamento = form_data.parcelamento.meio_pagamento
            if parcelamento_existente.valor_total is not None:
                parcelamento_existente.valor_total = form_data.parcelamento.valor_total
            if parcelamento_existente.valor_parcela is not None:
                parcelamento_existente.valor_parcela = form_data.parcelamento.valor_parcela
            if parcelamento_existente.valor_entrada is not None:
                parcelamento_existente.valor_entrada = form_data.parcelamento.valor_entrada
            if parcelamento_existente.qtd_parcela is not None:
                parcelamento_existente.qtd_parcela = form_data.parcelamento.qtd_parcela
            if parcelamento_existente.avista is not None:
                parcelamento_existente.avista = form_data.parcelamento.avista
            if parcelamento_existente.taxa_juros is not None:
                parcelamento_existente.taxa_juros = form_data.parcelamento.taxa_juros
            if parcelamento_existente.data_ultimo_pagamento is not None:
                parcelamento_existente.data_ultimo_pagamento = form_data.parcelamento.data_ultimo_pagamento
            if parcelamento_existente.qtd_parcelas_pagas is not None:
                parcelamento_existente.qtd_parcelas_pagas = form_data.parcelamento.qtd_parcelas_pagas
            parcelamento_existente.updated_by = uuid.UUID(user_id)
            parcelamento_existente.updated_at = datetime.utcnow()

            await db.merge(parcelamento_existente)
    else:
        novo_parcelamento_id = uuid.uuid4()
        novo_parcelamento = ParcelamentoModel(
            id=novo_parcelamento_id,
            contrato_id=form_data.id,
            data_inicio=form_data.parcelamento.data_inicio,
            data_fim=form_data.parcelamento.data_fim,
            data_vigencia=form_data.parcelamento.data_vigencia,
            meio_pagamento=form_data.parcelamento.meio_pagamento,
            valor_total=form_data.parcelamento.valor_total,
            valor_parcela=form_data.parcelamento.valor_parcela,
            valor_entrada=form_data.parcelamento.valor_entrada,
            qtd_parcela=form_data.parcelamento.qtd_parcela,
            avista=form_data.parcelamento.avista,
            taxa_juros=form_data.parcelamento.taxa_juros,
            data_ultimo_pagamento=form_data.parcelamento.data_ultimo_pagamento,
            qtd_parcelas_pagas=form_data.parcelamento.qtd_parcelas_pagas,
            ativo=True,
            created_by=uuid.UUID(user_id)
        )
        db.add(novo_parcelamento)
        await db.flush()
        parcelamento_id = novo_parcelamento_id

    contrato_existente.updated_at = datetime.utcnow()
    contrato_existente.parcelamento_id = parcelamento_id
    contrato_existente.cliente_id = form_data.cliente_id
    contrato_existente.vendedor_id = form_data.vendedor_id
    contrato_existente.plano_id = form_data.plano_id
    contrato_existente.anexos_list_id = anexos_final_ids
    contrato_existente.nome = cliente.nome
    contrato_existente.documento = cliente.documento
    contrato_existente.updated_by = uuid.UUID(user_id)

    dados_depois = limpar_dict_para_json(form_data)

    log = LogModel(
        tabela_afetada="contratos",
        operacao="UPDATE",
        registro_id=form_data.id,
        dados_antes=dados_antes,
        dados_depois=dados_depois,
        usuario_id=uuid.UUID(user_id)
    )
    db.add(log)

    await db.commit()
    await db.refresh(contrato_existente)

    return JSONResponse(
        content={"detail": "Contrato atualizado com sucesso"},
        media_type="application/json; charset=utf-8"
    )


async def delete_item(id: UUID, db: AsyncSession = Depends(get_db), user_id: str = Depends(verificar_token)):

    query = select(ContratoModel).where(ContratoModel.id == id)
    result = await db.execute(query)
    contrato = result.scalar_one_or_none()

    if not contrato:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contrato não encontrado."
        )

    dados_antigos = limpar_dict_para_json(contrato)

    contrato.ativo = False
    contrato.deleted_at = datetime.utcnow()
    contrato.deleted_by = uuid.UUID(user_id)

    dados_novos = limpar_dict_para_json(contrato)

    log = LogModel(
        tabela_afetada="contratos",
        operacao="DELETE",
        registro_id=contrato.id,
        dados_antes=dados_antigos,
        dados_depois=dados_novos,
        usuario_id=uuid.UUID(user_id)
    )

    db.add(log)
    await db.commit()

    return JSONResponse(
        content={"detail": "Contrato deletado com sucesso"},
        media_type="application/json; charset=utf-8"
    )


def valor_por_extenso(valor):
    reais = int(valor)
    centavos = int(round((valor - reais) * 100))
    extenso_reais = num2words(reais, lang='pt_BR')
    if centavos > 0:
        extenso_centavos = num2words(centavos, lang='pt_BR')
        return f"{extenso_reais} reais e {extenso_centavos} centavos"
    else:
        return f"{extenso_reais} reais"

def safe_get(obj, attr, default=""):
    return getattr(obj, attr, default) if obj else default

def get_data_format(data_inicio_raw):
    data_inicio = None
    if isinstance(data_inicio_raw, str):
        data_inicio = datetime.fromisoformat(data_inicio_raw.replace("Z", "+00:00"))
    else:
        data_inicio = data_inicio_raw

    return data_inicio

def contrato_pdf_pf(conteudo, styles, res):
    conteudo.append(Paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS FINANCEIROS PARA PESSOA FISICAS", styles['Titulo']))
    conteudo.append(Paragraph(
        f"Pelo presente instrumento, M D LIMA CONSULTORIA EIRELI,"
        f" inscrita no CNPJ:41.649.122/0001-90,"
        f" com sede social a Rua Doze de Outubro, nº 385 conjunto 23 – Lapa – CEP: 05073-001 – São Paulo – SP,"
        f" neste ato representado"
        f" MAYARA DANTAS LIMA, brasileira, solteira,"
        f" empresária, portadora do"
        f" CPF sob o nº. 106.592.994-37,"
        f" RG inscrito sob nº.59.879.242-9 SSP/SP,"
        f" doravante denominada MUTUANTE,"
        f" {safe_get(res.cliente, 'nome')}, telefone nº {safe_get(res.cliente, 'telefone')},"
        f" e-mail: {safe_get(res.cliente, 'email')} ,"
        f" inscrito sob o documento de nº {safe_get(res.cliente, 'documento')}, residente e domiciliado na "
        f" {safe_get(res.cliente.endereco, 'rua')}, nº {safe_get(res.cliente.endereco, 'numero')} {safe_get(res.cliente.endereco, 'complemento')},"
        f" {safe_get(res.cliente.endereco, 'bairro')}, CEP: {safe_get(res.cliente.endereco, 'cep')} e"
        f" {safe_get(res.cliente.endereco, 'cidade')}/{safe_get(res.cliente.endereco, 'uf')},"
        f" doravante denominado(a) MUTUÁRIO(A),"
        f" o presente mútuo, contrato nº {safe_get(res, 'numero')} mediante as seguintes cláusulas:",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA PRIMEIRA - DO OBJETO</b>", styles['Subtitulo']))

    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    valor_total = safe_get(res.parcelamento, 'valor_total')
    valor_total_formatado = locale.currency(valor_total, grouping=True)
    conteudo.append(Paragraph(
        f"1.1. Por meio do presente instrumento, o(a) MUTUANTE empresta ao(à) MUTUÁRIO (A), direta e pessoalmente, a quantia de {valor_total_formatado}"
        f" ({valor_por_extenso(valor_total)}). ",
        styles['Corpo']
    ))

    if res.parcelamento.qtd_parcela > 1:
        conteudo.append(Paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x.",
            styles['Corpo']
        ))
    else:
        conteudo.append(Paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio uma única parcela.",
            styles['Corpo']
        ))

    conteudo.append(Paragraph(
        f"1.3. O(A) MUTUANTE entregará a quantia ao(à) MUTUÁRIO(A) no ato de assinatura deste instrumento OU em"
        f" {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")}. ",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SEGUNDA - DA DESTINAÇÃO DO EMPRÉSTIMO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        f"O(A) MUTUÁRIO(A) poderá fazer livre uso da quantia emprestada, desde que não seja para fins econômicos, ou seja, fica vedada a alienação do valor.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA TERCEIRA - DO PAGAMENTO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        f"3.1.1 O(a) MUTUÁRIO(A) se compromete a restituir ao(à) MUTUANTE a quantia mutuada especificada na cláusula primeira, da seguinte forma:",
        styles['Corpo']
    ))

    if res.parcelamento.qtd_parcela > 1:
        conteudo.append(Paragraph(
            f"Parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x",
            styles['Corpo']
        ))
        datas_parcelas = []

        data_inicio = get_data_format(res.parcelamento.data_inicio)

        qtd_parcelas = int(safe_get(res.parcelamento, 'qtd_parcela'))

        for i in range(qtd_parcelas):
            data_parcela = data_inicio + relativedelta(months=i)
            datas_parcelas.append({
                "parcela": i + 1,
                "dia": data_parcela.strftime("%d/%m/%Y")
            })

        for data in datas_parcelas:
            conteudo.append(Paragraph(
                f"•	{data['parcela']}ª {data['dia']} ",
                styles['Corpo']
            ))
    else:
        conteudo.append(Paragraph(
            f"Parcela única",
            styles['Corpo']
        ))

        conteudo.append(Paragraph(
            f"•	1ª {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")} ",
            styles['Corpo']
        ))

    conteudo.append(Paragraph(
        "3.1.2. O empréstimo é realizado a título oneroso e haverá, portanto, incidência de juros compensatórios ou de "
        "correção monetário sobre o valor mutuado.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        f"3.1.3. O(A) MUTUÁRIO(A) se compromete a restituir o valor mutuado ao(à) MUTUANTE acrescido de juros de {safe_get(res.parcelamento, 'taxa_juros')}% "
        f"(teto máximo, podendo ser modificado) ao mês, aplicadas sobre a quantia total emprestada, além de correção monetária "
        f"calculada com base na variação do IGP-M do período.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.4. O atraso, bem como, o não pagamento fará com que o(a) MUTUANTE incorra em mora, sujeitando-se desta forma à "
        "cobranças extrajudiciais, bem como realização de protestos e o que se fizerem necessárias, com incidência de juros de 1% a.m. e de multa "
        "de 10% calculados sobre o mês de atraso.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.5 O MUTUÁRIO poderá amortizar ou liquidar a dívida do empréstimo, antes do vencimento.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.6. Eventual aceitação do(a) MUTUANTE(A) em receber parcelas pagas intempestivamente ou pelo não "
        "cumprimento de obrigações contratuais, a seu critério, não importará em novação, perdão ou alteração contratual, "
        "mas mera liberalidade do (a) MUTUÁRIO(A), permanecendo inalteradas as cláusulas deste contrato.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.7. O valor do débito, objeto deste contrato, ficará representado por uma Nota Promissória, emitida pelo MUTUANTE a favor do "
        "MUTUÁRIO, com vencimento à vista, avalizada pelo INTERVENIENTE, acima qualificado.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.8. O INTERVENIENTE, que é o avalista da supramencionada Nota Promissória, assinará este "
        "contrato também na qualidade de devedor solidário, no que atina ao pagamento da dívida contraída em razão deste instrumento.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.9 Se houver inadimplemento por parte do MUTUÁRIO(A), o MUTUANTE ficará autorizado a protestar ou executar a Nota Promissória, "
        "pelo valor do saldo devedor, apurado na época, e a executar a garantia real.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "PARAGRAFO ÚNICO - o MUTUÁRIO(A) desde já, autoriza que a(s) cobrança(s) seja(m) realizada em seu endereço residencial/comercial, "
        "desde que sejam respeitados os horários noturnos (que compreendem das 19:00hs ás 06:00 hs), ou aquele que o DEVEDOR indicar, "
        "inclusive aos finais de semana e feriados.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA QUARTA - DAS OBRIGAÇÕES</b>", styles['Subtitulo']))

    conteudo.append(Paragraph(
        "4.1. São obrigações do(a) MUTUÁRIO(a):",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "4.2. São obrigações do (a) MUTUANTE (A):",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Receber o pagamento da dívida, nos termos estipulados neste termo;",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Entregar recibo de quitação da dívida ao MUTUÁRIO(A), quando finalizado todo o pagamento previsto.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA QUINTA - DA CESSÃO E TRANSFERÊNCIA</b>", styles['Subtitulo']))

    conteudo.append(Paragraph(
        "5.1. Fica vedada a cessão e transferência do presente contrato, seja a que título for, sem a expressa "
        "concordância do MUTUANTE, havendo concordância, será realizado um novo contrato em nome do novo MUTUÁRIO.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SEXTA - DA SUCESSÃO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "6.1 Todas as obrigações assumidas neste instrumento são irrevogáveis e irretratáveis, o qual as partes obrigam-se "
        "a cumpri-lo, a qualquer título, e, em caso de óbito ou extinção de alguma das partes, serão transferidas a seus "
        "herdeiros ou sucessores, mediante anuência dos herdeiros.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SETIMA - DA VIGÊNCIA</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "7.1. O presente contrato passa a vigorar entre as partes a partir da assinatura dele.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA OITAVA - DO FORO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "8.1. As partes contratantes elegem o foro da cidade de São Paulo/SP para dirimir quaisquer dúvidas "
        "relativas ao cumprimento deste instrumento, não superadas pela mediação administrativa.",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "E, por estarem justos e combinados, MUTUANTE(A) e MUTUÁRIO(A) celebram e assinam o presente instrumento, "
        "em 2 (duas) vias de igual teor e forma, na presença das testemunhas, abaixo nomeadas e indicadas, que também "
        "o subscrevem, para que surta seus efeitos jurídicos.",
        styles['Corpo']
    ))

    conteudo.append(Spacer(1, 40))
    conteudo.append(Paragraph(f"São Paulo, {date.today().strftime('%d/%m/%Y')}", styles['Corpo']))

    assinaturas = [
        ["MUTUÁRIO", "MUTUANTE", "INTERVENIENTE"],
        ["__________________________", "__________________________", "__________________________"],
        [safe_get(res.cliente, 'nome'), "M D LIMA CONSULTORIA EIRELI ", "INTERVENIENTE"],
        [safe_get(res.cliente, 'documento'), "CNPJ: 41.649.122/0001-90 ", "INTERVENIENTE"]
    ]
    tabela = Table(assinaturas, colWidths=[250, 250])
    tabela.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 20),
    ]))
    conteudo.append(tabela)

    return conteudo

def contrato_pdf_pj(conteudo, styles, res):
    conteudo.append(Paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS FINANCEIROS PARA PESSOA JURÍDICA", styles['Titulo']))
    conteudo.append(Paragraph(
        f"Pelo presente instrumento, M D LIMA CONSULTORIA EIRELI,"
        f" inscrita no CNPJ:41.649.122/0001-90,"
        f" com sede social a Rua Doze de Outubro, nº 385 conjunto 23 – Lapa – CEP: 05073-001 – São Paulo – SP,"
        f" neste ato representado"
        f" MAYARA DANTAS LIMA, brasileira, solteira,"
        f" empresária, portadora do"
        f" CPF sob o nº. 106.592.994-37,"
        f" RG inscrito sob nº.59.879.242-9 SSP/SP,"
        f" doravante denominada MUTUANTE,"
        f" {safe_get(res.cliente, 'nome')},inscrito no CNPJ: {safe_get(res.cliente, 'documento')},"
        f" telefone nº {safe_get(res.cliente, 'telefone')},"
        f" e-mail: {safe_get(res.cliente, 'email')}, residente e domiciliado na "
        f" com sede social a Rua: "
        f" {safe_get(res.cliente.endereco, 'rua')}, nº {safe_get(res.cliente.endereco, 'numero')} {safe_get(res.cliente.endereco, 'complemento')},"
        f" {safe_get(res.cliente.endereco, 'bairro')}, CEP: {safe_get(res.cliente.endereco, 'cep')} e"
        f" {safe_get(res.cliente.endereco, 'cidade')}/{safe_get(res.cliente.endereco, 'uf')},"
        f" doravante denominado(a) MUTUÁRIO(A),"
        f" o presente mútuo, contrato nº {safe_get(res, 'numero')} mediante as seguintes cláusulas:",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA PRIMEIRA - DO OBJETO</b>", styles['Subtitulo']))

    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')  # define para português do Brasil

    valor_total = safe_get(res.parcelamento, 'valor_total')
    valor_total_formatado = locale.currency(valor_total, grouping=True)
    conteudo.append(Paragraph(
        f"1.1. Por meio do presente instrumento, o(a) MUTUANTE empresta ao(à) MUTUÁRIO (A), direta e pessoalmente, a quantia de {valor_total_formatado}"
        f" ({valor_por_extenso(valor_total)}). ",
        styles['Corpo']
    ))

    if res.parcelamento.qtd_parcela > 1:
        conteudo.append(Paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x.",
            styles['Corpo']
        ))
    else:
        conteudo.append(Paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio uma única parcela.",
            styles['Corpo']
        ))

    conteudo.append(Paragraph(
        f"1.3. O(A) MUTUANTE entregará a quantia ao(à) MUTUÁRIO(A) no ato de assinatura deste instrumento OU em"
        f" {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")}. ",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SEGUNDA - DA DESTINAÇÃO DO EMPRÉSTIMO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        f"O(A) MUTUÁRIO(A) poderá fazer livre uso da quantia emprestada, desde que não seja para fins econômicos, ou seja, fica vedada a alienação do valor.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA TERCEIRA - DO PAGAMENTO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        f"3.1.1 O(a) MUTUÁRIO(A) se compromete a restituir ao(à) MUTUANTE a quantia mutuada especificada na cláusula primeira, da seguinte forma:",
        styles['Corpo']
    ))

    if res.parcelamento.qtd_parcela > 1:
        conteudo.append(Paragraph(
            f"Parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x",
            styles['Corpo']
        ))
        datas_parcelas = []

        data_inicio = get_data_format(res.parcelamento.data_inicio)

        qtd_parcelas = int(safe_get(res.parcelamento, 'qtd_parcela'))

        for i in range(qtd_parcelas):
            data_parcela = data_inicio + relativedelta(months=i)
            datas_parcelas.append({
                "parcela": i + 1,
                "dia": data_parcela.strftime("%d/%m/%Y")
            })

        for data in datas_parcelas:
            conteudo.append(Paragraph(
                f"•	{data['parcela']}ª {data['dia']} ",
                styles['Corpo']
            ))
    else:
        conteudo.append(Paragraph(
            f"Parcela única",
            styles['Corpo']
        ))

        conteudo.append(Paragraph(
            f"•	1ª {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")} ",
            styles['Corpo']
        ))

    conteudo.append(Paragraph(
        "3.1.2. O empréstimo é realizado a título oneroso e haverá, portanto, incidência de juros compensatórios ou de "
        "correção monetário sobre o valor mutuado.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        f"3.1.3. O(A) MUTUÁRIO(A) se compromete a restituir o valor mutuado ao(à) MUTUANTE acrescido de juros de {safe_get(res.parcelamento, 'taxa_juros')}% "
        f"(teto máximo, podendo ser modificado) ao mês, aplicadas sobre a quantia total emprestada, além de correção monetária "
        f"calculada com base na variação do IGP-M do período.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.4. O atraso, bem como, o não pagamento fará com que o(a) MUTUANTE incorra em mora, sujeitando-se desta forma à "
        "cobranças extrajudiciais, bem como realização de protestos e o que se fizerem necessárias, com incidência de juros de 1% a.m. e de multa "
        "de 10% calculados sobre o mês de atraso.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.5 O MUTUÁRIO poderá amortizar ou liquidar a dívida do empréstimo, antes do vencimento.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.6. Eventual aceitação do(a) MUTUANTE(A) em receber parcelas pagas intempestivamente ou pelo não "
        "cumprimento de obrigações contratuais, a seu critério, não importará em novação, perdão ou alteração contratual, "
        "mas mera liberalidade do (a) MUTUÁRIO(A), permanecendo inalteradas as cláusulas deste contrato.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.7. O valor do débito, objeto deste contrato, ficará representado por uma Nota Promissória, emitida pelo MUTUANTE a favor do "
        "MUTUÁRIO, com vencimento à vista, avalizada pelo INTERVENIENTE, acima qualificado.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.8. O INTERVENIENTE, que é o avalista da supramencionada Nota Promissória, assinará este "
        "contrato também na qualidade de devedor solidário, no que atina ao pagamento da dívida contraída em razão deste instrumento.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "3.1.9 Se houver inadimplemento por parte do MUTUÁRIO(A), o MUTUANTE ficará autorizado a protestar ou executar a Nota Promissória, "
        "pelo valor do saldo devedor, apurado na época, e a executar a garantia real.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "PARAGRAFO ÚNICO - o MUTUÁRIO(A) desde já, autoriza que a(s) cobrança(s) seja(m) realizada em seu endereço residencial/comercial, "
        "desde que sejam respeitados os horários noturnos (que compreendem das 19:00hs ás 06:00 hs), ou aquele que o DEVEDOR indicar, "
        "inclusive aos finais de semana e feriados.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA QUARTA - DAS OBRIGAÇÕES</b>", styles['Subtitulo']))

    conteudo.append(Paragraph(
        "4.1. São obrigações do(a) MUTUÁRIO(a):",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;",
        styles['Corpo']
    ))

    conteudo.append(Paragraph(
        "4.2. São obrigações do (a) MUTUANTE (A):",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Receber o pagamento da dívida, nos termos estipulados neste termo;",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "•	Entregar recibo de quitação da dívida ao MUTUÁRIO(A), quando finalizado todo o pagamento previsto.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA QUINTA - DA CESSÃO E TRANSFERÊNCIA</b>", styles['Subtitulo']))

    conteudo.append(Paragraph(
        "5.1. Fica vedada a cessão e transferência do presente contrato, seja a que título for, sem a expressa "
        "concordância do MUTUANTE, havendo concordância, será realizado um novo contrato em nome do novo MUTUÁRIO.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SEXTA - DA SUCESSÃO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "6.1 Todas as obrigações assumidas neste instrumento são irrevogáveis e irretratáveis, o qual as partes obrigam-se "
        "a cumpri-lo, a qualquer título, e, em caso de óbito ou extinção de alguma das partes, serão transferidas a seus "
        "herdeiros ou sucessores, mediante anuência dos herdeiros.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA SETIMA - DA VIGÊNCIA</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "7.1. O presente contrato passa a vigorar entre as partes a partir da assinatura dele.",
        styles['Corpo']
    ))

    conteudo.append(Paragraph("<b>CLÁUSULA OITAVA - DO FORO</b>", styles['Subtitulo']))
    conteudo.append(Paragraph(
        "8.1. As partes contratantes elegem o foro da cidade de São Paulo/SP para dirimir quaisquer dúvidas "
        "relativas ao cumprimento deste instrumento, não superadas pela mediação administrativa.",
        styles['Corpo']
    ))
    conteudo.append(Paragraph(
        "E, por estarem justos e combinados, MUTUANTE(A) e MUTUÁRIO(A) celebram e assinam o presente instrumento, "
        "em 2 (duas) vias de igual teor e forma, na presença das testemunhas, abaixo nomeadas e indicadas, que também "
        "o subscrevem, para que surta seus efeitos jurídicos.",
        styles['Corpo']
    ))

    conteudo.append(Spacer(1, 40))
    conteudo.append(Paragraph(f"São Paulo, {date.today().strftime('%d/%m/%Y')}", styles['Corpo']))

    assinaturas = [
        ["MUTUÁRIO", "MUTUANTE", "INTERVENIENTE"],
        ["__________________________", "__________________________", "__________________________"],
        [safe_get(res.cliente, 'nome'), "M D LIMA CONSULTORIA EIRELI ", "INTERVENIENTE"],
        [safe_get(res.cliente, 'documento'), "CNPJ: 41.649.122/0001-90 ", "INTERVENIENTE"]
    ]
    tabela = Table(assinaturas, colWidths=[250, 250])
    tabela.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 20),
    ]))
    conteudo.append(tabela)

    return conteudo

async def gerar_contrato_pdf(id: UUID, db: AsyncSession = Depends(get_db), user_id: str = Depends(verificar_token)):
    res = await por_id(id, db)

    if not res:
        raise HTTPException(status_code=400, detail="Contrato não localizado na base de dados.")

    buffer = BytesIO()

    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40, leftMargin=40,
        topMargin=60, bottomMargin=60
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Titulo',
                              fontSize=20,
                              leading=24,
                              alignment=1,
                              textColor=colors.HexColor("#003366"),
                              spaceAfter=20))
    styles.add(ParagraphStyle(name='Subtitulo',
                              fontSize=14,
                              leading=18,
                              textColor=colors.HexColor("#005599"),
                              spaceAfter=10))
    styles.add(ParagraphStyle(name='Corpo',
                              fontSize=12,
                              leading=16,
                              alignment=4,
                              spaceAfter=10))

    conteudo = []

    doc = ''.join(filter(str.isdigit, str(res.cliente.documento)))

    if len(doc) == 11:
        conteudo = contrato_pdf_pf(conteudo, styles, res)
    else:
        conteudo = contrato_pdf_pj(conteudo, styles, res)

    pdf.build(conteudo)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=contrato.pdf"}
    )


def contrato_word_pf(doc, res):
    doc.add_paragraph(
        f"Pelo presente instrumento, M D LIMA CONSULTORIA EIRELI,"
        f" inscrita no CNPJ:41.649.122/0001-90,"
        f" com sede social a Rua Doze de Outubro, nº 385 conjunto 23 – Lapa – CEP: 05073-001 – São Paulo – SP,"
        f" neste ato representado"
        f" MAYARA DANTAS LIMA, brasileira, solteira,"
        f" empresária, portadora do"
        f" CPF sob o nº. 106.592.994-37,"
        f" RG inscrito sob nº.59.879.242-9 SSP/SP,"
        f" doravante denominada MUTUANTE,"
        f" {safe_get(res.cliente, 'nome')}, telefone nº {safe_get(res.cliente, 'telefone')},"
        f" e-mail: {safe_get(res.cliente, 'email')} ,"
        f" inscrito sob o documento de nº {safe_get(res.cliente, 'documento')}, residente e domiciliado na "
        f" {safe_get(res.cliente.endereco, 'rua')}, nº {safe_get(res.cliente.endereco, 'numero')} {safe_get(res.cliente.endereco, 'complemento')},"
        f" {safe_get(res.cliente.endereco, 'bairro')}, CEP: {safe_get(res.cliente.endereco, 'cep')} e"
        f" {safe_get(res.cliente.endereco, 'cidade')}/{safe_get(res.cliente.endereco, 'uf')},"
        f" doravante denominado(a) MUTUÁRIO(A),"
        f" o presente mútuo, contrato nº {safe_get(res, 'numero')} mediante as seguintes cláusulas:",
    )
    doc.add_heading("CLÁUSULA PRIMEIRA - DO OBJETO", level=2)

    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    valor_total = safe_get(res.parcelamento, 'valor_total')
    valor_total_formatado = locale.currency(valor_total, grouping=True)

    doc.add_paragraph(
        f"1.1. Por meio do presente instrumento, o(a) MUTUANTE empresta ao(à) MUTUÁRIO (A), direta e pessoalmente, a quantia de {valor_total_formatado}"
        f" ({valor_por_extenso(valor_total)}). ",
    )

    if res.parcelamento.qtd_parcela > 1:
        doc.add_paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x."
        )
    else:
        doc.add_paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio uma única parcela.")

    doc.add_paragraph(
        f"1.3. O(A) MUTUANTE entregará a quantia ao(à) MUTUÁRIO(A) no ato de assinatura deste instrumento OU em"
        f" {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")}. "
    )

    doc.add_heading("CLÁUSULA SEGUNDA - DA DESTINAÇÃO DO EMPRÉSTIMO", level=2)
    doc.add_paragraph(
        "O(A) MUTUÁRIO(A) poderá fazer livre uso da quantia emprestada, desde que não seja para fins econômicos, ou seja, fica vedada a alienação do valor."
    )

    doc.add_heading("CLÁUSULA TERCEIRA - DO PAGAMENTO", level=2)
    doc.add_paragraph(
        f"3.1.1 O(a) MUTUÁRIO(A) se compromete a restituir ao(à) MUTUANTE a quantia mutuada especificada na cláusula primeira, da seguinte forma:"
    )

    if res.parcelamento.qtd_parcela > 1:
        doc.add_paragraph(
            f"Parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x"
        )
        datas_parcelas = []

        data_inicio = get_data_format(res.parcelamento.data_inicio)

        qtd_parcelas = int(safe_get(res.parcelamento, 'qtd_parcela'))

        for i in range(qtd_parcelas):
            data_parcela = data_inicio + relativedelta(months=i)
            datas_parcelas.append({
                "parcela": i + 1,
                "dia": data_parcela.strftime("%d/%m/%Y")
            })

        for data in datas_parcelas:
            doc.add_paragraph(
                f"•	{data['parcela']}ª {data['dia']} "
            )
    else:
        doc.add_paragraph(
            f"Parcela única"
        )

        doc.add_paragraph(
            f"•	1ª {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")} "
        )

    doc.add_paragraph(
        "3.1.2. O empréstimo é realizado a título oneroso e haverá, portanto, incidência de juros compensatórios ou de "
        "correção monetário sobre o valor mutuado."
    )
    doc.add_paragraph(
        f"3.1.3. O(A) MUTUÁRIO(A) se compromete a restituir o valor mutuado ao(à) MUTUANTE acrescido de juros de {safe_get(res.parcelamento, 'taxa_juros')}% "
        f"(teto máximo, podendo ser modificado) ao mês, aplicadas sobre a quantia total emprestada, além de correção monetária "
        f"calculada com base na variação do IGP-M do período."
    )

    doc.add_paragraph(
        "3.1.4. O atraso, bem como, o não pagamento fará com que o(a) MUTUANTE incorra em mora, sujeitando-se desta forma à cobranças extrajudiciais, bem como realização de protestos e o que se fizerem necessárias, com incidência de juros de 1% a.m. e de multa de 10% calculados sobre o mês de atraso."
    )

    doc.add_paragraph(
        "3.1.5 O MUTUÁRIO poderá amortizar ou liquidar a dívida do empréstimo, antes do vencimento."
    )

    doc.add_paragraph(
        "3.1.6. Eventual aceitação do(a) MUTUANTE(A) em receber parcelas pagas intempestivamente ou pelo não cumprimento de obrigações contratuais, a seu critério, não importará em novação, perdão ou alteração contratual, mas mera liberalidade do (a) MUTUÁRIO(A), permanecendo inalteradas as cláusulas deste contrato."
    )

    doc.add_paragraph(
        "3.1.7. O valor do débito, objeto deste contrato, ficará representado por uma Nota Promissória, emitida pelo MUTUANTE a favor do MUTUÁRIO, com vencimento à vista, avalizada pelo INTERVENIENTE, acima qualificado."
    )

    doc.add_paragraph(
        "3.1.8. O INTERVENIENTE, que é o avalista da supramencionada Nota Promissória, assinará este contrato também na qualidade de devedor solidário, no que atina ao pagamento da dívida contraída em razão deste instrumento."
    )

    doc.add_paragraph(
        "3.1.9 Se houver inadimplemento por parte do MUTUÁRIO(A), o MUTUANTE ficará autorizado a protestar ou executar a Nota Promissória, pelo valor do saldo devedor, apurado na época, e a executar a garantia real."
    )

    doc.add_paragraph(
        "PARAGRAFO ÚNICO - o MUTUÁRIO(A) desde já, autoriza que a(s) cobrança(s) seja(m) realizada em seu endereço residencial/comercial, desde que sejam respeitados os horários noturnos (que compreendem das 19:00hs ás 06:00 hs), ou aquele que o DEVEDOR indicar, inclusive aos finais de semana e feriados."
    )

    doc.add_heading("CLÁUSULA QUARTA - DAS OBRIGAÇÕES", level=2)

    doc.add_paragraph(
        "4.1. São obrigações do(a) MUTUÁRIO(a):"
    )
    doc.add_paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;"
    )
    doc.add_paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;"
    )

    doc.add_paragraph(
        "4.2. São obrigações do (a) MUTUANTE (A):"
    )
    doc.add_paragraph(
        "•	Receber o pagamento da dívida, nos termos estipulados neste termo;"
    )
    doc.add_paragraph(
        "•	Entregar recibo de quitação da dívida ao MUTUÁRIO(A), quando finalizado todo o pagamento previsto."
    )

    doc.add_heading("CLÁUSULA QUINTA - DA CESSÃO E TRANSFERÊNCIA", level=2)
    doc.add_paragraph(
        "5.1. Fica vedada a cessão e transferência do presente contrato, seja a que título for, sem a expressa concordância do MUTUANTE, havendo concordância, será realizado um novo contrato em nome do novo MUTUÁRIO."
    )

    doc.add_heading("CLÁUSULA SEXTA - DA SUCESSÃO ", level=2)
    doc.add_paragraph(
        "6.1 Todas as obrigações assumidas neste instrumento são irrevogáveis e irretratáveis, o qual as partes obrigam-se a cumpri-lo, a qualquer título, e, em caso de óbito ou extinção de alguma das partes, serão transferidas a seus herdeiros ou sucessores, mediante anuência dos herdeiros."
    )

    doc.add_heading("CLÁUSULA SETIMA - DA VIGÊNCIA", level=2)
    doc.add_paragraph(
        "7.1. O presente contrato passa a vigorar entre as partes a partir da assinatura dele."
    )

    doc.add_heading("CLÁUSULA OITAVA - DO FORO", level=2)
    doc.add_paragraph(
        "8.1. As partes contratantes elegem o foro da cidade de São Paulo/SP para dirimir quaisquer dúvidas relativas ao cumprimento deste instrumento, não superadas pela mediação administrativa."
    )
    doc.add_paragraph(
        "E, por estarem justos e combinados, MUTUANTE(A) e MUTUÁRIO(A) celebram e assinam o presente instrumento, em 2 (duas) vias de igual teor e forma, na presença das testemunhas, abaixo nomeadas e indicadas, que também o subscrevem, para que surta seus efeitos jurídicos."
    )

    # Data e assinaturas
    doc.add_paragraph(f"\nSão Paulo, {date.today().strftime('%d/%m/%Y')}.")

    doc.add_paragraph("MUTUÁRIO")
    doc.add_paragraph(f"\n______________________________\n{safe_get(res.cliente, 'nome')}\n{safe_get(res.cliente, 'documento')}\n")
    doc.add_paragraph(f"\n")
    doc.add_paragraph("MUTUANTE")
    doc.add_paragraph("\n______________________________\nM D LIMA CONSULTORIA EIRELI\nCNPJ: 41.649.122/0001-90\n")
    doc.add_paragraph(f"\n")
    doc.add_paragraph("INTERVENIENTE")
    doc.add_paragraph("\n______________________________\nINTERVENIENTE\n")
    doc.add_paragraph(f"\n")

    return doc


def contrato_word_pj(doc, res):
    doc.add_paragraph(
        f"Pelo presente instrumento, M D LIMA CONSULTORIA EIRELI,"
        f" inscrita no CNPJ:41.649.122/0001-90,"
        f" com sede social a Rua Doze de Outubro, nº 385 conjunto 23 – Lapa – CEP: 05073-001 – São Paulo – SP,"
        f" neste ato representado"
        f" MAYARA DANTAS LIMA, brasileira, solteira,"
        f" empresária, portadora do"
        f" CPF sob o nº. 106.592.994-37,"
        f" RG inscrito sob nº.59.879.242-9 SSP/SP,"
        f" doravante denominada MUTUANTE,"
        f" {safe_get(res.cliente, 'nome')},inscrito no CNPJ: {safe_get(res.cliente, 'documento')},"
        f" telefone nº {safe_get(res.cliente, 'telefone')},"
        f" e-mail: {safe_get(res.cliente, 'email')}, residente e domiciliado na "
        f" com sede social a Rua: "
        f" {safe_get(res.cliente.endereco, 'rua')}, nº {safe_get(res.cliente.endereco, 'numero')} {safe_get(res.cliente.endereco, 'complemento')},"
        f" {safe_get(res.cliente.endereco, 'bairro')}, CEP: {safe_get(res.cliente.endereco, 'cep')} e"
        f" {safe_get(res.cliente.endereco, 'cidade')}/{safe_get(res.cliente.endereco, 'uf')},"
        f" doravante denominado(a) MUTUÁRIO(A),"
        f" o presente mútuo, contrato nº {safe_get(res, 'numero')} mediante as seguintes cláusulas:",
    )
    doc.add_heading("CLÁUSULA PRIMEIRA - DO OBJETO", level=2)

    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    valor_total = safe_get(res.parcelamento, 'valor_total')
    valor_total_formatado = locale.currency(valor_total, grouping=True)

    doc.add_paragraph(
        f"1.1. Por meio do presente instrumento, o(a) MUTUANTE empresta ao(à) MUTUÁRIO (A), direta e pessoalmente, a quantia de {valor_total_formatado}"
        f" ({valor_por_extenso(valor_total)}). ",
    )

    if res.parcelamento.qtd_parcela > 1:
        doc.add_paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x."
        )
    else:
        doc.add_paragraph(
            f"1.2. A quantia será repassada ao(à) MUTUÁRIO(A) mediante, por meio uma única parcela.")

    doc.add_paragraph(
        f"1.3. O(A) MUTUANTE entregará a quantia ao(à) MUTUÁRIO(A) no ato de assinatura deste instrumento OU em"
        f" {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")}. "
    )

    doc.add_heading("CLÁUSULA SEGUNDA - DA DESTINAÇÃO DO EMPRÉSTIMO", level=2)
    doc.add_paragraph(
        "O(A) MUTUÁRIO(A) poderá fazer livre uso da quantia emprestada, desde que não seja para fins econômicos, ou seja, fica vedada a alienação do valor."
    )

    doc.add_heading("CLÁUSULA TERCEIRA - DO PAGAMENTO", level=2)
    doc.add_paragraph(
        f"3.1.1 O(a) MUTUÁRIO(A) se compromete a restituir ao(à) MUTUANTE a quantia mutuada especificada na cláusula primeira, da seguinte forma:"
    )

    if res.parcelamento.qtd_parcela > 1:
        doc.add_paragraph(
            f"Parcelado em {safe_get(res.parcelamento, 'qtd_parcela')}x"
        )
        datas_parcelas = []

        data_inicio = get_data_format(res.parcelamento.data_inicio)

        qtd_parcelas = int(safe_get(res.parcelamento, 'qtd_parcela'))

        for i in range(qtd_parcelas):
            data_parcela = data_inicio + relativedelta(months=i)
            datas_parcelas.append({
                "parcela": i + 1,
                "dia": data_parcela.strftime("%d/%m/%Y")
            })

        for data in datas_parcelas:
            doc.add_paragraph(
                f"•	{data['parcela']}ª {data['dia']} "
            )
    else:
        doc.add_paragraph(
            f"Parcela única"
        )

        doc.add_paragraph(
            f"•	1ª {get_data_format(res.parcelamento.data_inicio).strftime("%d/%m/%Y")} "
        )

    doc.add_paragraph(
        "3.1.2. O empréstimo é realizado a título oneroso e haverá, portanto, incidência de juros compensatórios ou de "
        "correção monetário sobre o valor mutuado."
    )
    doc.add_paragraph(
        f"3.1.3. O(A) MUTUÁRIO(A) se compromete a restituir o valor mutuado ao(à) MUTUANTE acrescido de juros de {safe_get(res.parcelamento, 'taxa_juros')}% "
        f"(teto máximo, podendo ser modificado) ao mês, aplicadas sobre a quantia total emprestada, além de correção monetária "
        f"calculada com base na variação do IGP-M do período."
    )

    doc.add_paragraph(
        "3.1.4. O atraso, bem como, o não pagamento fará com que o(a) MUTUANTE incorra em mora, sujeitando-se desta forma à cobranças extrajudiciais, bem como realização de protestos e o que se fizerem necessárias, com incidência de juros de 1% a.m. e de multa de 10% calculados sobre o mês de atraso."
    )

    doc.add_paragraph(
        "3.1.5 O MUTUÁRIO poderá amortizar ou liquidar a dívida do empréstimo, antes do vencimento."
    )

    doc.add_paragraph(
        "3.1.6. Eventual aceitação do(a) MUTUANTE(A) em receber parcelas pagas intempestivamente ou pelo não cumprimento de obrigações contratuais, a seu critério, não importará em novação, perdão ou alteração contratual, mas mera liberalidade do (a) MUTUÁRIO(A), permanecendo inalteradas as cláusulas deste contrato."
    )

    doc.add_paragraph(
        "3.1.7. O valor do débito, objeto deste contrato, ficará representado por uma Nota Promissória, emitida pelo MUTUANTE a favor do MUTUÁRIO, com vencimento à vista, avalizada pelo INTERVENIENTE, acima qualificado."
    )

    doc.add_paragraph(
        "3.1.8. O INTERVENIENTE, que é o avalista da supramencionada Nota Promissória, assinará este contrato também na qualidade de devedor solidário, no que atina ao pagamento da dívida contraída em razão deste instrumento."
    )

    doc.add_paragraph(
        "3.1.9 Se houver inadimplemento por parte do MUTUÁRIO(A), o MUTUANTE ficará autorizado a protestar ou executar a Nota Promissória, pelo valor do saldo devedor, apurado na época, e a executar a garantia real."
    )

    doc.add_paragraph(
        "PARAGRAFO ÚNICO - o MUTUÁRIO(A) desde já, autoriza que a(s) cobrança(s) seja(m) realizada em seu endereço residencial/comercial, desde que sejam respeitados os horários noturnos (que compreendem das 19:00hs ás 06:00 hs), ou aquele que o DEVEDOR indicar, inclusive aos finais de semana e feriados."
    )

    doc.add_heading("CLÁUSULA QUARTA - DAS OBRIGAÇÕES", level=2)

    doc.add_paragraph(
        "4.1. São obrigações do(a) MUTUÁRIO(a):"
    )
    doc.add_paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;"
    )
    doc.add_paragraph(
        "•	Efetuar o pagamento pontualmente, conforme as datas e os meios fixados neste instrumento;"
    )

    doc.add_paragraph(
        "4.2. São obrigações do (a) MUTUANTE (A):"
    )
    doc.add_paragraph(
        "•	Receber o pagamento da dívida, nos termos estipulados neste termo;"
    )
    doc.add_paragraph(
        "•	Entregar recibo de quitação da dívida ao MUTUÁRIO(A), quando finalizado todo o pagamento previsto."
    )

    doc.add_heading("CLÁUSULA QUINTA - DA CESSÃO E TRANSFERÊNCIA", level=2)
    doc.add_paragraph(
        "5.1. Fica vedada a cessão e transferência do presente contrato, seja a que título for, sem a expressa concordância do MUTUANTE, havendo concordância, será realizado um novo contrato em nome do novo MUTUÁRIO."
    )

    doc.add_heading("CLÁUSULA SEXTA - DA SUCESSÃO ", level=2)
    doc.add_paragraph(
        "6.1 Todas as obrigações assumidas neste instrumento são irrevogáveis e irretratáveis, o qual as partes obrigam-se a cumpri-lo, a qualquer título, e, em caso de óbito ou extinção de alguma das partes, serão transferidas a seus herdeiros ou sucessores, mediante anuência dos herdeiros."
    )

    doc.add_heading("CLÁUSULA SETIMA - DA VIGÊNCIA", level=2)
    doc.add_paragraph(
        "7.1. O presente contrato passa a vigorar entre as partes a partir da assinatura dele."
    )

    doc.add_heading("CLÁUSULA OITAVA - DO FORO", level=2)
    doc.add_paragraph(
        "8.1. As partes contratantes elegem o foro da cidade de São Paulo/SP para dirimir quaisquer dúvidas relativas ao cumprimento deste instrumento, não superadas pela mediação administrativa."
    )
    doc.add_paragraph(
        "E, por estarem justos e combinados, MUTUANTE(A) e MUTUÁRIO(A) celebram e assinam o presente instrumento, em 2 (duas) vias de igual teor e forma, na presença das testemunhas, abaixo nomeadas e indicadas, que também o subscrevem, para que surta seus efeitos jurídicos."
    )

    # Data e assinaturas
    doc.add_paragraph(f"\nSão Paulo, {date.today().strftime('%d/%m/%Y')}.")

    doc.add_paragraph("MUTUÁRIO")
    doc.add_paragraph(
        f"\n______________________________\n{safe_get(res.cliente, 'nome')}\n{safe_get(res.cliente, 'documento')}\n")
    doc.add_paragraph(f"\n")
    doc.add_paragraph("MUTUANTE")
    doc.add_paragraph("\n______________________________\nM D LIMA CONSULTORIA EIRELI\nCNPJ: 41.649.122/0001-90\n")
    doc.add_paragraph(f"\n")
    doc.add_paragraph("INTERVENIENTE")
    doc.add_paragraph("\n______________________________\nINTERVENIENTE\n")
    doc.add_paragraph(f"\n")

    return doc


async def gerar_contrato_word(id: UUID, db: AsyncSession = Depends(get_db), user_id: str = Depends(verificar_token)):
    res = await por_id(id, db)
    if not res:
        raise HTTPException(status_code=400, detail="Contrato não localizado na base de dados.")

    doc = Document()

    titulo = doc.add_paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS FINANCEIROS PARA PESSOA FISICAS")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].bold = True
    titulo.runs[0].font.size = Pt(14)


    documento = ''.join(filter(str.isdigit, str(res.cliente.documento)))

    if len(documento) == 11:
        doc = contrato_word_pf(doc, res)
    else:
        doc = contrato_word_pj(doc, res)

    nome_arquivo = "contrato.docx"
    doc.save(nome_arquivo)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=contrato.docx"}
    )